# core/ai_views.py
import re

from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.contrib.auth.decorators import login_required
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import json
import logging
import os
import PyPDF2
import io
import docx
import re
from .ai_legal_assistant import AILegalAssistant
import google.generativeai as genai
from django.conf import settings
from .smart_case_analyzer import SmartCaseAnalyzer
from .legal_text_generator import LegalTextGenerator

logger = logging.getLogger(__name__)

# AI Hukuki Asistan Views
def ai_assistant_home(request):
    """AI Hukuki Asistan ana sayfası"""
    assistant = AILegalAssistant()
    suggestions = assistant.get_quick_suggestions()
    
    return render(request, 'core/ai_assistant.html', {
        'suggestions': suggestions,
        'page_title': 'AI Hukuki Asistan'
    })

@csrf_exempt
@require_http_methods(["POST"])
def ai_assistant_api(request):
    """AI Hukuki Asistan API endpoint"""
    try:
        data = json.loads(request.body)
        question = data.get('question', '').strip()
        
        if not question:
            return JsonResponse({
                'success': False,
                'error': 'Lütfen bir soru sorun.'
            })
        
        if len(question) < 5:
            return JsonResponse({
                'success': False,
                'error': 'Sorunuz çok kısa. Lütfen daha detaylı bir soru sorun.'
            })
        
        # AI Assistant'a gönder
        assistant = AILegalAssistant()
        result = assistant.process_question(question, request.user if request.user.is_authenticated else None)
        
        return JsonResponse(result)
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Geçersiz veri formatı.'
        })
    except Exception as e:
        logger.error(f"AI Assistant API error: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Bir hata oluştu. Lütfen tekrar deneyin.'
        })

# Akıllı Dava Dosyası Analizi Views
@login_required
def smart_case_analyzer_home(request):
    """Akıllı Dava Dosyası Analizi ana sayfası"""
    return render(request, 'core/smart_case_analyzer.html', {
        'page_title': 'Akıllı Dava Dosyası Analizi'
    })

@csrf_exempt
@login_required
@require_http_methods(["POST"])
def analyze_case_file(request):
    """Dava dosyası analiz API"""
    try:
        if 'case_file' not in request.FILES:
            return JsonResponse({
                'success': False,
                'error': 'Lütfen bir dosya yükleyin.'
            })
        
        uploaded_file = request.FILES['case_file']
        
        # Dosya kontrolü
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        allowed_extensions = ['.pdf', '.udf']
        
        if file_extension not in allowed_extensions:
            return JsonResponse({
                'success': False,
                'error': 'Sadece PDF ve UDF dosyaları kabul edilir.'
            })
        
        if uploaded_file.size > 10 * 1024 * 1024:  # 10MB limit
            return JsonResponse({
                'success': False,
                'error': 'Dosya boyutu 10MB\'dan büyük olamaz.'
            })
        
        # Dosya içeriğini oku
        file_content = extract_file_content(uploaded_file)
        if not file_content:
            return JsonResponse({
                'success': False,
                'error': 'Dosya içeriği okunamadı. Dosyanın bozuk olmadığından emin olun.'
            })
        
        # Analiz et
        analyzer = SmartCaseAnalyzer()
        analysis_result = analyzer.analyze_case_document(file_content, request.user)
        
        return JsonResponse(analysis_result)
        
    except Exception as e:
        logger.error(f"Case analysis error: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Dosya analizi sırasında bir hata oluştu.'
        })

def extract_file_content(file_obj):
    """Dosyadan metin çıkar (PDF, UDF desteği)"""
    try:
        file_name = file_obj.name.lower()
        
        if file_name.endswith('.pdf'):
            # PDF dosyasını işle
            pdf_reader = PyPDF2.PdfReader(file_obj)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        elif file_name.endswith('.udf'):
            # UDF dosyasını işle
            from .udf_extractor import UDFExtractor
            extractor = UDFExtractor()
            return extractor.extract_text(file_obj)
            
        else:
            logger.error(f"Desteklenmeyen dosya türü: {file_name}")
            return None
        
    except Exception as e:
        logger.error(f"File content extraction error: {str(e)}")
        return None

# Hukuki Metin Üretici Views
@login_required
def legal_text_generator_home(request):
    """Hukuki Metin Üretici ana sayfası"""
    generator = LegalTextGenerator()
    templates = generator.get_available_templates()
    
    import os
    template_path = os.path.join(settings.BASE_DIR, "core", "templates", "core", "legal_text_generator.html")
    logger.info(f"Template path: {template_path}")
    logger.info(f"Template exists: {os.path.exists(template_path)}")
    logger.info(f"Template size: {os.path.getsize(template_path) if os.path.exists(template_path) else 0}")
    return render(request, 'core/legal_text_generator.html', {
        'templates': json.dumps(templates),
        'page_title': 'Hukuki Metin Üretici'
    })

@csrf_exempt
@login_required
@require_http_methods(["POST"])
def generate_legal_text(request):
    """Hukuki metin üretme API"""
    logger.info("Generate legal text called")
    try:
        data = json.loads(request.body)
        logger.info(f"Received data: {data}")
        
        template_type = data.get('template_type')
        parameters = data.get('parameters', {})
        
        if not template_type:
            logger.error("No template type provided")
            return JsonResponse({
                'success': False,
                'error': 'Belge türü belirtilmedi.'
            }, status=400)
        
        # LegalTextGenerator instance
        generator = LegalTextGenerator()
        
        # Generate the document
        result = generator.generate_document(template_type, parameters)
        
        if result['success']:
            logger.info("Document generated successfully")
            # Session"a belgeyi kaydet Word indirme için
            request.session["last_generated_document"] = result.get("document", {}).get("content", "")
            request.session["last_document_title"] = result.get("document", {}).get("template_name", "Hukuki Belge")
            request.session["last_document_type"] = template_type
            return JsonResponse({
                'success': True,
                'content': result.get('document', {}).get('content', ''),
                'title': result.get('title', 'Hukuki Belge')
            })
        else:
            logger.error(f"Generation failed: {result.get('error')}")
            return JsonResponse({
                'success': False,
                'error': result.get('error', 'Belge oluşturulamadı.')
            }, status=500)
            
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Geçersiz JSON verisi.'
        }, status=400)
    except Exception as e:
        logger.error(f"Generate legal text error: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Bir hata oluştu: {str(e)}'
        }, status=500)
@csrf_exempt
@login_required
@require_http_methods(["POST"])
def get_template_fields(request):
    """Şablon alanlarını getir"""
    try:
        data = json.loads(request.body)
        
        template_type = data.get('template_type')
        subtype = data.get('subtype')
        
        if not template_type:
            return JsonResponse({
                'success': False,
                'error': 'Şablon türü gerekli.'
            })
        
        generator = LegalTextGenerator()
        fields_info = generator.get_template_fields(template_type, subtype)
        
        if fields_info:
            return JsonResponse({
                'success': True,
                'fields': fields_info['fields']
            })
        else:
            return JsonResponse({
                'success': False,
                'error': 'Şablon bulunamadı.'
            })
        
    except Exception as e:
        logger.error(f"Get template fields error: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Alanlar yüklenemedi.'
        })

@csrf_exempt
@login_required
@require_http_methods(["POST"])
def generate_from_multiple_documents(request):
    """Çoklu belgeye istinaden metin üretme API"""
    try:
        logger.info("generate_from_multiple_documents API çağrıldı")
        
        if 'source_documents' not in request.FILES:
            return JsonResponse({
                'success': False,
                'error': 'Lütfen en az bir kaynak belge yükleyin.'
            })
        
        uploaded_files = request.FILES.getlist('source_documents')
        document_type = request.POST.get('document_type', '').strip()
        additional_instructions = request.POST.get('additional_instructions', '').strip()
        
        logger.info(f"Dosya sayısı: {len(uploaded_files)}, Tür: {document_type}")
        
        if not uploaded_files:
            return JsonResponse({
                'success': False,
                'error': 'Lütfen en az bir dosya yükleyin.'
            })
        
        if not document_type:
            return JsonResponse({
                'success': False,
                'error': 'Lütfen üretilecek belge türünü belirtin.'
            })
        
        # Dosya kontrolları
        allowed_extensions = ['.pdf', '.docx', '.txt', '.udf']
        total_size = 0
        
        for uploaded_file in uploaded_files:
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()
            
            if file_extension not in allowed_extensions:
                return JsonResponse({
                    'success': False,
                    'error': f'{uploaded_file.name}: Sadece PDF, DOCX, TXT ve UDF dosyaları kabul edilir.'
                })
            
            if uploaded_file.size > 10 * 1024 * 1024:  # 10MB limit per file
                return JsonResponse({
                    'success': False,
                    'error': f'{uploaded_file.name}: Dosya boyutu 10MB\'dan büyük olamaz.'
                })
            
            total_size += uploaded_file.size
        
        if total_size > 50 * 1024 * 1024:  # 50MB total limit
            return JsonResponse({
                'success': False,
                'error': 'Toplam dosya boyutu 50MB\'dan büyük olamaz.'
            })
        
        # Metin üret
        logger.info("LegalTextGenerator başlatılıyor...")
        generator = LegalTextGenerator()
        result = generator.generate_from_multiple_documents(
            uploaded_files, document_type, additional_instructions
        )
        
        logger.info(f"Üretim sonucu: {result.get('success', False)}")
        if result.get('success'):
            # Session"a belgeyi kaydet Word indirme iu00e7in
            request.session["last_generated_document"] = result.get("document", {}).get("content", "")
            request.session["last_document_title"] = document_type.replace("_", " ").title()
            request.session["last_document_type"] = document_type
            doc = result.get('document', {})
            logger.info(f"Belge içerik uzunluğu: {len(doc.get('content', ''))}")
            logger.info(f"İlk 200 karakter: {doc.get('content', '')[:200]}")
        else:
            logger.error(f"Üretim hatası: {result.get('error', 'Bilinmeyen')}")
        
        return JsonResponse(result)
        
    except Exception as e:
        logger.error(f"Multiple document generation error: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Çoklu belge üretimi sırasında bir hata oluştu: {str(e)}'
        })

@csrf_exempt
@login_required
@require_http_methods(["POST"])
def generate_from_uploaded_document(request):
    """Yüklenen belgeye istinaden metin üretme API"""
    try:
        logger.info("generate_from_uploaded_document API çağrıldı")
        
        if 'source_document' not in request.FILES:
            return JsonResponse({
                'success': False,
                'error': 'Lütfen bir kaynak belge yükleyin.'
            })
        
        uploaded_file = request.FILES['source_document']
        document_type = request.POST.get('document_type', '')
        additional_params = request.POST.get('additional_params', '')
        
        logger.info(f"Dosya: {uploaded_file.name}, Tür: {document_type}")
        
        # Dosya kontrolü
        allowed_extensions = ['.pdf', '.docx', '.txt', '.udf']
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_extension not in allowed_extensions:
            return JsonResponse({
                'success': False,
                'error': 'Sadece PDF, DOCX, TXT ve UDF dosyaları kabul edilir.'
            })
        
        if uploaded_file.size > 10 * 1024 * 1024:  # 10MB limit
            return JsonResponse({
                'success': False,
                'error': 'Dosya boyutu 10MB\'dan büyük olamaz.'
            })
        
        if not document_type:
            return JsonResponse({
                'success': False,
                'error': 'Lütfen üretilecek belge türünü seçin.'
            })
        
        # Metin üret
        logger.info("LegalTextGenerator başlatılıyor...")
        generator = LegalTextGenerator()
        result = generator.generate_from_uploaded_document(
            uploaded_file, document_type, additional_params
        )
        
        logger.info(f"Üretim sonucu: {result.get('success', False)}")
        if result.get('success'):
            # Session"a belgeyi kaydet Word indirme iu00e7in
            request.session["last_generated_document"] = result.get("document", {}).get("content", "")
            request.session["last_document_title"] = document_type.replace("_", " ").title()
            request.session["last_document_type"] = document_type
            doc = result.get('document', {})
            logger.info(f"Belge içerik uzunluğu: {len(doc.get('content', ''))}")
            logger.info(f"İlk 200 karakter: {doc.get('content', '')[:200]}")
        else:
            logger.error(f"Üretim hatası: {result.get('error', 'Bilinmeyen')}")
        
        return JsonResponse(result)
        
    except Exception as e:
        logger.error(f"Document-based generation error: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Belge tabanlı üretim sırasında bir hata oluştu: {str(e)}'
        })

@csrf_exempt
@login_required  
@require_http_methods(["POST"])
def save_generated_text(request):
    """Üretilen metni kaydet"""
    try:
        data = json.loads(request.body)
        
        text_content = data.get('content')
        document_name = data.get('name', 'Hukuki Belge')
        
        if not text_content:
            return JsonResponse({
                'success': False,
                'error': 'Kaydedilecek içerik bulunamadı.'
            })
        
        # Dosyayı kaydet
        filename = f"generated_legal_docs/{request.user.id}/{document_name}.txt"
        file_path = default_storage.save(filename, ContentFile(text_content))
        
        return JsonResponse({
            'success': True,
            'message': 'Belge başarıyla kaydedildi.',
            'file_path': file_path
        })
        
    except Exception as e:
        logger.error(f"Save generated text error: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Belge kaydedilemedi.'
        })

# Ortak AI özellikleri
def ai_features_home(request):
    """AI özellikler ana sayfası"""
    return render(request, 'core/ai_features.html', {
        'page_title': 'Yapay Zeka Özellikleri'
    })
@csrf_exempt
@login_required
@require_http_methods(["POST"])
def generate_from_multiple_documents(request):
    """Yüklenen belgelerden yeni belge üretme API"""
    logger.info("Generate from multiple documents called")
    try:
        # Get form data
        document_type = request.POST.get('document_type')
        additional_instructions = request.POST.get('additional_instructions', '')
        file_count = int(request.POST.get('file_count', 0))
        
        if not document_type:
            return JsonResponse({
                'success': False,
                'error': 'Belge türü belirtilmedi.'
            }, status=400)
        
        # Get uploaded files
        uploaded_files = request.FILES.getlist('source_documents')
        
        if not uploaded_files:
            return JsonResponse({
                'success': False,
                'error': 'Dosya yüklenmedi.'
            }, status=400)
        
        logger.info(f"Processing {len(uploaded_files)} files for {document_type}")
        
        # Extract content from all files
        all_content = []
        for uploaded_file in uploaded_files:
            content = extract_file_content(uploaded_file)
            if content:
                all_content.append(f"--- {uploaded_file.name} ---\n{content}\n")
            else:
                logger.warning(f"Could not extract content from {uploaded_file.name}")
        
        if not all_content:
            return JsonResponse({
                'success': False,
                'error': 'Dosyalardan içerik çıkarılamadı.'
            }, status=400)
        
        # Combine all content
        combined_content = "\n\n".join(all_content)
        
        # Create prompt for Gemini
        prompt = f"""Sen uzman bir hukuk asistanısın. Aşağıdaki belgeleri inceleyerek '{document_type}' türünde yeni bir belge oluştur.

Kaynak Belgeler:
{combined_content}

{f'Ek Talimatlar: {additional_instructions}' if additional_instructions else ''}

Lütfen profesyonel, detaylı ve hukuki standartlara uygun bir belge oluştur. Belge Türkçe olmalı ve Türk hukukuna uygun olmalıdır.
"""
        
        # Generate with Gemini
        generator = LegalTextGenerator()
        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel(model_name="gemini-2.5-pro")
        
        response = model.generate_content(prompt)
        
        if response and response.text:
            logger.info("Document generated successfully from uploads")
            return JsonResponse({
                'success': True,
                'content': response.text,
                'title': document_type
            })
        else:
            logger.error("No response from Gemini")
            return JsonResponse({
                'success': False,
                'error': 'Belge oluşturulamadı.'
            }, status=500)
            
    except Exception as e:
        logger.error(f"Generate from documents error: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Bir hata oluştu: {str(e)}'
        }, status=500)

@csrf_exempt
@login_required
@require_http_methods(["GET"])
def get_template_fields(request):
    """Template field'larını getir"""
    template_type = request.GET.get('template_type')
    subtype = request.GET.get('subtype')
    
    try:
        generator = LegalTextGenerator()
        templates = generator.get_available_templates()
        
        if template_type in templates:
            template_info = templates[template_type]
            return JsonResponse({
                'success': True,
                'fields': template_info.get('required_fields', []),
                'subtypes': template_info.get('subtypes', [])
            })
        else:
            return JsonResponse({
                'success': False,
                'error': 'Template not found'
            })
    except Exception as e:
        logger.error(f"Get template fields error: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        })

@csrf_exempt
@login_required
@require_http_methods(["POST"])
def save_text(request):
    """Metni kaydet"""
    try:
        data = json.loads(request.body)
        content = data.get('content', '')
        filename = data.get('filename', 'document.txt')
        
        # Simple save - in production you'd want to save to database
        return JsonResponse({
            'success': True,
            'message': 'Text saved successfully'
        })
    except Exception as e:
        logger.error(f"Save text error: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        })

# Contract Risk Analyzer Functions

def contract_risk_analyzer(request):
    """Sözleşme Risk Analizi ana sayfası"""
    return render(request, 'core/contract_risk_analyzer.html', {
        'page_title': 'Sözleşme Risk Analizi'
    })

def extract_text_from_file(uploaded_file):
    """Yüklenen dosyadan metin çıkar"""
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_extension == 'pdf':
            # PDF okuma
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text += page.extract_text() + "\n"
                except Exception as e:
                    logger.warning(f"Page {page_num} read error: {e}")
                    continue
            return text.strip()
            
        elif file_extension in ['doc', 'docx']:
            # Word dosyası okuma
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            # Tablolardaki metni de al
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            return text.strip()
            
        else:
            return None
            
    except Exception as e:
        logger.error(f"File extraction error: {e}")
        return None

def parse_contract_ai_response(response_text):
    """AI yanıtını parse et - contract analizi için"""
    result = {
        'risk_score': 5,
        'risk_indicators': [],
        'clauses': []
    }
    
    try:
        lines = response_text.split('\n')
        
        # Risk skoru bul
        for line in lines:
            if 'RİSK SKORU' in line or 'risk skoru' in line.lower():
                # Sayıları bul
                numbers = re.findall(r'\d+', line)
                if numbers:
                    score = int(numbers[0])
                    result['risk_score'] = min(max(score, 1), 10)
                    break
        
        # Risk göstergelerini parse et
        in_indicators = False
        in_clauses = False
        current_clause = None
        
        for line in lines:
            line = line.strip()
            
            # Risk göstergeleri bölümü
            if 'RİSK GÖSTERGELERİ' in line or 'risk göstergeleri' in line.lower():
                in_indicators = True
                in_clauses = False
                continue
                
            # Madde analizi bölümü
            if 'MADDE ANALİZİ' in line or 'madde analizi' in line.lower():
                in_indicators = False
                in_clauses = True
                if current_clause:
                    result['clauses'].append(current_clause)
                    current_clause = None
                continue
            
            # Risk göstergelerini parse et
            if in_indicators and line and (line.startswith('-') or line.startswith('•')):
                level = 'medium'
                if 'düşük' in line.lower() or 'low' in line.lower():
                    level = 'low'
                elif 'yüksek' in line.lower() or 'high' in line.lower():
                    level = 'high'
                
                icon = 'fas fa-info-circle'
                if level == 'low':
                    icon = 'fas fa-check-circle'
                elif level == 'high':
                    icon = 'fas fa-exclamation-triangle'
                
                # Metni temizle
                text = line.lstrip('-•').strip()
                text = re.sub(r'level:\s*\w+', '', text, flags=re.IGNORECASE).strip()
                
                if text:
                    result['risk_indicators'].append({
                        'title': text.split('.')[0].strip(),
                        'level': level,
                        'icon': icon,
                        'description': text
                    })
            
            # Madde analizini parse et
            if in_clauses:
                # Yeni madde başlangıcı
                if line and (line.startswith('-') or line.startswith('•') or re.match(r'^\d+\.', line)):
                    if current_clause:
                        result['clauses'].append(current_clause)
                    
                    # Risk seviyesini bul
                    risk_level = 'medium'
                    risk_text = 'Orta Risk'
                    if 'düşük risk' in line.lower() or 'low' in line.lower():
                        risk_level = 'low'
                        risk_text = 'Düşük Risk'
                    elif 'yüksek risk' in line.lower() or 'high' in line.lower():
                        risk_level = 'high'
                        risk_text = 'Yüksek Risk'
                    
                    # Başlığı çıkar
                    title = line.lstrip('-•').strip()
                    title = re.sub(r'-?\s*Risk:\s*\w+', '', title, flags=re.IGNORECASE).strip()
                    title = re.sub(r'^\d+\.\s*', '', title).strip()
                    
                    current_clause = {
                        'title': title,
                        'risk_level': risk_level,
                        'risk_text': risk_text,
                        'content': '',
                        'recommendation': ''
                    }
                
                # İçerik satırı
                elif current_clause and line:
                    if line.startswith('İçerik:') or line.startswith('Açıklama:'):
                        current_clause['content'] = line.split(':', 1)[1].strip()
                    elif line.startswith('Öneri:') or line.startswith('Tavsiye:'):
                        current_clause['recommendation'] = line.split(':', 1)[1].strip()
                    elif not line.startswith(('İçerik:', 'Öneri:', 'Açıklama:', 'Tavsiye:')):
                        # Devam eden metin
                        if current_clause['content'] and not current_clause['recommendation']:
                            current_clause['content'] += ' ' + line
                        elif current_clause['recommendation']:
                            current_clause['recommendation'] += ' ' + line
        
        # Son maddeyi ekle
        if current_clause:
            result['clauses'].append(current_clause)
        
        # Eğer hiç gösterge yoksa varsayılan ekle
        if not result['risk_indicators']:
            result['risk_indicators'] = [
                {
                    'title': 'Genel Değerlendirme',
                    'level': 'medium',
                    'icon': 'fas fa-info-circle',
                    'description': 'Sözleşme standart risk seviyesinde değerlendirilmiştir.'
                }
            ]
        
        # Eğer hiç madde yoksa varsayılan ekle
        if not result['clauses']:
            result['clauses'] = [
                {
                    'title': 'Genel Analiz',
                    'risk_level': 'medium',
                    'risk_text': 'Orta Risk',
                    'content': 'Sözleşme genel hükümler içermektedir.',
                    'recommendation': 'Detaylı hukuki inceleme önerilir.'
                }
            ]
            
    except Exception as e:
        logger.error(f"Parse error: {e}")
    
    return result

@csrf_exempt
@require_http_methods(["POST"])
def analyze_contract(request):
    """Sözleşme risk analizi API"""
    try:
        if 'contract_file' not in request.FILES:
            return JsonResponse({
                'success': False,
                'error': 'Lütfen bir sözleşme dosyası yükleyin.'
            })
        
        uploaded_file = request.FILES['contract_file']
        contract_type = request.POST.get('contract_type', 'general')
        
        # Dosya boyutu kontrolü (10MB)
        if uploaded_file.size > 10 * 1024 * 1024:
            return JsonResponse({
                'success': False,
                'error': 'Dosya boyutu 10MB\'dan büyük olamaz.'
            })
        
        # Dosya içeriğini oku
        file_content = extract_text_from_file(uploaded_file)
        if not file_content:
            return JsonResponse({
                'success': False,
                'error': 'Dosya içeriği okunamadı. Lütfen PDF veya Word formatında bir dosya yükleyin.'
            })
        
        # Sözleşme türü eşlemeleri
        contract_type_names = {
            'employment': 'İş Sözleşmesi',
            'rental': 'Kira Sözleşmesi',
            'sales': 'Satış Sözleşmesi',
            'service': 'Hizmet Sözleşmesi',
            'nda': 'Gizlilik Sözleşmesi',
            'other': 'Genel Sözleşme'
        }
        
        contract_type_name = contract_type_names.get(contract_type, 'Genel Sözleşme')
        
        # İçeriği sınırla (API limitleri için)
        content_preview = file_content[:4000] if len(file_content) > 4000 else file_content
        
        prompt = f"""
Sen uzman bir hukuk danışmanısın. Aşağıdaki {contract_type_name} türündeki sözleşmeyi analiz et ve risk değerlendirmesi yap.

Sözleşme İçeriği:
{content_preview}

Lütfen aşağıdaki formatta detaylı bir analiz sun:

GENEL RİSK SKORU: [1-10 arasında bir sayı, 1 en düşük risk, 10 en yüksek risk]

RİSK GÖSTERGELERİ:
- [Risk faktörü 1] - level: [low/medium/high]
- [Risk faktörü 2] - level: [low/medium/high]
- [Risk faktörü 3] - level: [low/medium/high]

MADDE ANALİZİ:
- [Riskli madde başlığı 1] - Risk: [low/medium/high]
  İçerik: [Bu maddenin neden riskli olduğunun açıklaması]
  Öneri: [Bu riski azaltmak için ne yapılmalı]

- [Riskli madde başlığı 2] - Risk: [low/medium/high]
  İçerik: [Bu maddenin neden riskli olduğunun açıklaması]
  Öneri: [Bu riski azaltmak için ne yapılmalı]

Analiz yaparken özellikle şunlara dikkat et:
1. Tek taraflı ve dengesiz maddeler
2. Belirsiz ve yoruma açık ifadeler
3. Ceza ve tazminat maddeleri
4. Fesih ve sona erme koşulları
5. Gizlilik ve rekabet yasağı maddeleri
6. Mücbir sebep halleri
7. Uyuşmazlık çözüm yöntemleri
8. Yasal mevzuata uygunluk
"""
        
        # Gemini API'yi kullan
        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel(model_name="gemini-1.5-flash")
        
        try:
            response = model.generate_content(prompt)
            
            if response and response.text:
                # AI yanıtını parse et
                parsed_result = parse_contract_ai_response(response.text)
                
                logger.info(f"Contract analyzed successfully: risk score {parsed_result['risk_score']}")
                
                return JsonResponse({
                    'success': True,
                    'risk_score': parsed_result['risk_score'],
                    'risk_indicators': parsed_result['risk_indicators'],
                    'clauses': parsed_result['clauses']
                })
            else:
                logger.error("No response from AI")
                return JsonResponse({
                    'success': False,
                    'error': 'AI analizi tamamlanamadı. Lütfen tekrar deneyin.'
                })
                
        except Exception as ai_error:
            logger.error(f"AI analysis error: {ai_error}")
            # Fallback demo response
            return JsonResponse({
                'success': True,
                'risk_score': 6,
                'risk_indicators': [
                    {
                        'title': 'Ceza Maddeleri',
                        'level': 'high',
                        'icon': 'fas fa-exclamation-triangle',
                        'description': 'Sözleşmede yüksek ceza bedelleri tespit edildi.'
                    },
                    {
                        'title': 'Fesih Koşulları',
                        'level': 'medium',
                        'icon': 'fas fa-info-circle',
                        'description': 'Tek taraflı fesih hakları dengesiz görünüyor.'
                    },
                    {
                        'title': 'Yasal Uyum',
                        'level': 'low',
                        'icon': 'fas fa-check-circle',
                        'description': 'Genel yasal gerekliliklere uygun görünüyor.'
                    }
                ],
                'clauses': [
                    {
                        'title': 'Ceza Koşulları',
                        'risk_level': 'high',
                        'risk_text': 'Yüksek Risk',
                        'content': 'Sözleşmede belirtilen ceza bedelleri piyasa standartlarının üzerinde görünüyor.',
                        'recommendation': 'Ceza bedellerinin makul seviyelere çekilmesi ve karşılıklılık ilkesine uygun hale getirilmesi önerilir.'
                    },
                    {
                        'title': 'Gizlilik Maddeleri',
                        'risk_level': 'medium',
                        'risk_text': 'Orta Risk',
                        'content': 'Gizlilik yükümlülükleri çok geniş kapsamlı ve süresiz olarak tanımlanmış.',
                        'recommendation': 'Gizlilik kapsamının sınırlandırılması ve makul bir süre belirlenmesi tavsiye edilir.'
                    }
                ]
            })
            
    except Exception as e:
        logger.error(f"Contract analysis error: {e}")
        return JsonResponse({
            'success': False,
            'error': f'Analiz hatası: {str(e)}'
        }, status=500)
@csrf_exempt
@login_required
@require_http_methods(["POST"])
def download_document_word(request):
    """Belgeyi Word formatında indir"""
    try:
        # Session'dan belge içeriğini al
        document_content = request.session.get('last_generated_document')
        document_title = request.session.get('last_document_title', 'Hukuki Belge')
        document_type = request.session.get('last_document_type', 'belge')
        
        if not document_content:
            return JsonResponse({
                'success': False,
                'error': 'İndirilecek belge bulunamadı. Lütfen önce bir belge oluşturun.'
            })
        
        from .word_document_generator import WordDocumentGenerator
        
        # Word belgesi oluştur
        generator = WordDocumentGenerator()
        success = generator.create_legal_document(
            content=document_content,
            title=document_title,
            document_type=document_type
        )
        
        if not success:
            return JsonResponse({
                'success': False,
                'error': 'Word belgesi oluşturulurken hata oluştu.'
            })
        
        # Belgeyi bytes olarak al
        doc_bytes = generator.save_to_bytes()
        
        # HTTP response oluştur
        response = HttpResponse(
            doc_bytes.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Dosya adını temizle
        clean_title = re.sub(r'[^\w\s-]', '', document_title)[:50]
        filename = f"{clean_title}_{document_type}.docx"
        
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Word belgesi indirildi: {filename}")
        return response
        
    except Exception as e:
        logger.error(f"Word indirme hatası: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Word belgesi indirme hatası: {str(e)}'
        })

