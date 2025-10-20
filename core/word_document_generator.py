import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

class WordDocumentGenerator:
    """Word belgesi oluşturucu sınıfı"""
    
    def __init__(self):
        self.document = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Belge stillerini ayarla"""
        # Varsayılan font ayarları
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
    
    def clean_asterisks(self, text):
        """Metindeki markdown formatlamasını temizle"""
        if not text:
            return text
            
        # **bold** -> BÜYÜK HARF
        text = re.sub(r'\*\*(.*?)\*\*', lambda m: m.group(1).upper(), text)
        
        # *italic* -> Normal metin
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        
        # Tek * işaretlerini temizle
        text = text.replace('*', '')
        
        return text
    
    def create_legal_document(self, content, title="Hukuki Belge", document_type="belge"):
        """Hukuki belge oluştur"""
        try:
            # İçeriği temizle
            clean_content = self.clean_asterisks(content)
            
            # Başlık ekle
            title_paragraph = self.document.add_heading(title, 1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ana içerik
            paragraphs = clean_content.split('\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    p = self.document.add_paragraph()
                    
                    # Başlık kontrolü
                    if paragraph_text.strip().isupper() or paragraph_text.endswith(':'):
                        run = p.add_run(paragraph_text)
                        run.bold = True
                    else:
                        p.add_run(paragraph_text)
                    
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Alt kısım - tarih ve imza
            self.document.add_paragraph()
            date_p = self.document.add_paragraph()
            date_p.add_run("Tarih: .................................")
            date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            self.document.add_paragraph()
            signature_p = self.document.add_paragraph()
            signature_p.add_run("İmza: .................................")
            signature_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            logger.info(f"Word belgesi oluşturuldu: {title}")
            return True
            
        except Exception as e:
            logger.error(f"Word belgesi oluşturma hatası: {e}")
            return False
    
    def save_to_bytes(self):
        """Belgeyi bytes olarak döndür"""
        doc_bytes = BytesIO()
        self.document.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
    
    def save_to_file(self, filename):
        """Belgeyi dosyaya kaydet"""
        self.document.save(filename)
        return filename
